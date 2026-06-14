"""Document persistence: file storage, text extraction, and DB CRUD."""

from __future__ import annotations

import shutil
import uuid
from pathlib import Path

import docx
import pypdf
from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.config import settings
from app.db.models import Document
from app.utils.logger import get_logger

logger = get_logger(__name__)

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class UnsupportedFileTypeError(ValueError):
    """Raised when an uploaded file has an unsupported extension."""


class DocumentNotFoundError(LookupError):
    """Raised when a document id cannot be found."""


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #
def _extract_pdf(path: Path) -> str:
    reader = pypdf.PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    return "\n".join(parts).strip()


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def extract_text(path: Path, file_type: str) -> str:
    """Extract plain text from a stored file based on its extension."""
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
    }
    extractor = extractors.get(file_type.lower())
    if extractor is None:
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")
    text = extractor(path)
    if not text:
        logger.warning("No extractable text found in %s", path.name)
    return text


# --------------------------------------------------------------------------- #
# File storage
# --------------------------------------------------------------------------- #
def save_upload(file: UploadFile) -> tuple[str, Path, int, str]:
    """Persist an uploaded file to the uploads directory.

    Returns ``(stored_filename, path, size_bytes, extension)``.
    Raises ``UnsupportedFileTypeError`` for disallowed extensions.
    """
    original_name = file.filename or "upload"
    ext = Path(original_name).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Allowed: {sorted(ALLOWED_EXTENSIONS)}"
        )

    settings.ensure_dirs()
    stored_filename = f"{uuid.uuid4().hex}{ext}"
    dest = settings.uploads_dir / stored_filename

    with dest.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    size_bytes = dest.stat().st_size
    max_bytes = settings.max_upload_mb * 1024 * 1024
    if size_bytes > max_bytes:
        dest.unlink(missing_ok=True)
        raise UnsupportedFileTypeError(
            f"File too large ({size_bytes} bytes). Max is {settings.max_upload_mb} MB."
        )

    logger.info("Stored upload %s -> %s (%d bytes)", original_name, stored_filename, size_bytes)
    return stored_filename, dest, size_bytes, ext


# --------------------------------------------------------------------------- #
# Database CRUD
# --------------------------------------------------------------------------- #
def create_document(
    db: Session,
    *,
    owner_id: str,
    filename: str,
    stored_filename: str,
    file_path: Path,
    file_type: str,
    size_bytes: int,
    status: str = "pending",
) -> Document:
    doc = Document(
        owner_id=owner_id,
        filename=filename,
        stored_filename=stored_filename,
        file_path=str(file_path),
        file_type=file_type,
        size_bytes=size_bytes,
        status=status,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


def update_status(
    db: Session,
    doc: Document,
    *,
    status: str,
    num_chunks: int | None = None,
    error: str | None = None,
) -> Document:
    doc.status = status
    if num_chunks is not None:
        doc.num_chunks = num_chunks
    doc.error = error
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


def list_documents(db: Session, owner_id: str) -> list[Document]:
    return (
        db.query(Document)
        .filter(Document.owner_id == owner_id)
        .order_by(Document.created_at.desc())
        .all()
    )


def get_document(db: Session, document_id: str, owner_id: str) -> Document | None:
    doc = db.get(Document, document_id)
    # Scope to the owner: treat another owner's document as if it does not exist.
    if doc is None or doc.owner_id != owner_id:
        return None
    return doc


def require_document(db: Session, document_id: str, owner_id: str) -> Document:
    doc = get_document(db, document_id, owner_id)
    if doc is None:
        raise DocumentNotFoundError(f"Document not found: {document_id}")
    return doc


def load_text(db: Session, document_id: str, owner_id: str) -> tuple[Document, str]:
    """Return the owner's document record and its freshly extracted full text."""
    doc = require_document(db, document_id, owner_id)
    text = extract_text(Path(doc.file_path), doc.file_type)
    return doc, text
